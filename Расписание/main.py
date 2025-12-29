# -*- coding: utf-8 -*-

import sys
from docx import Document
from docx.oxml.ns import qn
from html import escape

def get_paragraph_alignment(paragraph):
    """Получить выравнивание параграфа"""
    if paragraph.alignment:
        alignment = paragraph.alignment
        if alignment == 0:  # LEFT
            return 'left'
        elif alignment == 1:  # CENTER
            return 'center'
        elif alignment == 2:  # RIGHT
            return 'right'
        elif alignment == 3:  # JUSTIFY
            return 'justify'
    return 'left'

def get_cell_alignment(cell):
    """Получить выравнивание ячейки таблицы"""
    # Проверяем выравнивание первого параграфа в ячейке
    if cell.paragraphs:
        return get_paragraph_alignment(cell.paragraphs[0])
    return 'left'

def get_cell_span_info(cell):
    """Получить информацию об объединении ячеек"""
    tc = cell._tc
    colspan = 1
    rowspan = 1
    
    # Проверяем горизонтальное объединение (colspan)
    grid_span = tc.find(qn('w:gridSpan'))
    if grid_span is not None:
        colspan = int(grid_span.get(qn('w:val'), 1))
    
    # Проверяем вертикальное объединение (rowspan)
    vmerge = tc.find(qn('w:vMerge'))
    if vmerge is not None:
        # Если это начало объединения
        if vmerge.get(qn('w:val')) is None or vmerge.get(qn('w:val')) == 'restart':
            # Подсчитываем количество объединенных строк
            rowspan = 1
            row_index = cell._parent._parent.index(cell._parent)
            table = cell._parent._parent._parent
            for i in range(row_index + 1, len(table.rows)):
                if len(table.rows[i].cells) > table.rows[row_index].cells.index(cell):
                    next_cell = table.rows[i].cells[table.rows[row_index].cells.index(cell)]
                    next_vmerge = next_cell._tc.find(qn('w:vMerge'))
                    if next_vmerge is not None and (next_vmerge.get(qn('w:val')) == 'continue' or next_vmerge.get(qn('w:val')) is None):
                        rowspan += 1
                    else:
                        break
                else:
                    break
        else:
            # Это продолжение объединения - пропускаем
            return 0, 0
    
    return colspan, rowspan

def process_paragraph(paragraph):
    """Обработать параграф и преобразовать в HTML"""
    alignment = get_paragraph_alignment(paragraph)
    style_name = paragraph.style.name.lower() if paragraph.style else ''
    
    # Определяем тег в зависимости от стиля
    if 'heading' in style_name:
        try:
            level = int(style_name.split()[-1])
            level = max(1, min(6, level))
            tag_name = f'h{level}'
        except (ValueError, IndexError):
            tag_name = 'h1'
    else:
        tag_name = 'p'
    
    # Добавляем стиль выравнивания
    style_attr = f' style="text-align: {alignment}"' if alignment != 'left' else ''
    
    html_parts = [f'<{tag_name}{style_attr}>']
    
    # Обрабатываем текстовые фрагменты
    for run in paragraph.runs:
        styles = []
        
        # Проверяем форматирование
        if run.bold:
            styles.append('font-weight: bold;')
        if run.italic:
            styles.append('font-style: italic;')
        if run.underline:
            styles.append('text-decoration: underline;')
        
        # Размер шрифта (уменьшен для мобильных устройств)
        if run.font.size and run.font.size.pt:
            # Уменьшаем размер шрифта для мобильных
            font_size_pt = max(8, run.font.size.pt * 0.7)  # Уменьшаем на 30%, но не менее 8pt
            styles.append(f'font-size: {font_size_pt}pt;')
        else:
            # Базовый размер шрифта для мобильных
            styles.append('font-size: 10pt;')
        
        # Цвет текста
        if run.font.color and run.font.color.rgb:
            rgb_hex = f'#{run.font.color.rgb:06x}'
            styles.append(f'color: {rgb_hex};')
        
        text = escape(run.text)
        
        if styles:
            style_str = ' '.join(styles)
            html_parts.append(f'<span style="{style_str}">{text}</span>')
        else:
            html_parts.append(text)
    
    html_parts.append(f'</{tag_name}>')
    
    return ''.join(html_parts)

def process_table(table):
    """Обработать таблицу и преобразовать в HTML"""
    html_parts = ['<div class="table-container">']
    html_parts.append('<table>')
    
    # Создаем матрицу для отслеживания объединенных ячеек
    processed_cells = set()
    
    for row_idx, row in enumerate(table.rows):
        html_parts.append('<tr>')
        
        col_idx = 0
        for cell in row.cells:
            # Пропускаем уже обработанные ячейки (объединенные)
            if (row_idx, col_idx) in processed_cells:
                col_idx += 1
                continue
            
            # Получаем информацию об объединении
            colspan, rowspan = get_cell_span_info(cell)
            
            # Пропускаем ячейки, которые являются продолжением объединения
            if colspan == 0 and rowspan == 0:
                col_idx += 1
                continue
            
            # Помечаем объединенные ячейки как обработанные
            for i in range(rowspan):
                for j in range(colspan):
                    if i == 0 and j == 0:
                        continue  # Основная ячейка уже будет обработана
                    processed_cells.add((row_idx + i, col_idx + j))
            
            # Формируем атрибуты ячейки
            cell_attrs = []
            if colspan > 1:
                cell_attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                cell_attrs.append(f'rowspan="{rowspan}"')
            
            # Добавляем выравнивание
            alignment = get_cell_alignment(cell)
            if alignment != 'left':
                cell_attrs.append(f'style="text-align: {alignment}"')
            
            # Добавляем классы для первой строки и первого столбца
            cell_classes = []
            if row_idx == 0:
                cell_classes.append('first-row')
            if col_idx == 0:
                cell_classes.append('first-col')
            
            if cell_classes:
                cell_attrs.append(f'class="{" ".join(cell_classes)}"')
            
            attr_str = ' '.join(cell_attrs)
            if attr_str:
                attr_str = ' ' + attr_str
            
            # Обрабатываем содержимое ячейки
            cell_content = []
            for paragraph in cell.paragraphs:
                if paragraph.text.strip() or paragraph.runs:  # Не пропускаем параграфы с форматированием
                    cell_content.append(process_paragraph(paragraph))
            
            content = '<br>'.join(cell_content) if cell_content else '&nbsp;'
            
            # Определяем тег ячейки (th для заголовков или td)
            cell_tag = 'th' if ((row_idx == 0 or col_idx == 0) and cell.paragraphs and 'heading' in str(cell.paragraphs[0].style.name).lower()) else 'td'
            
            html_parts.append(f'<{cell_tag}{attr_str}>{content}</{cell_tag}>')
            col_idx += colspan
        
        html_parts.append('</tr>')
    
    html_parts.append('</table>')
    html_parts.append('</div>')
    return '\n'.join(html_parts)

def docx_to_html(docx_path, html_path):
    """Преобразовать DOCX файл в HTML"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Ошибка при открытии файла: {e}")
        return False
    
    html_content = []
    html_content.append('<!DOCTYPE html>')
    html_content.append('<html>')
    html_content.append('<head>')
    html_content.append('<meta charset="UTF-8">')
    html_content.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    html_content.append('<title>Converted Document</title>')
    html_content.append('<style>')
    html_content.append('* { box-sizing: border-box; }')
    html_content.append('body { ')
    html_content.append('  font-family: Arial, sans-serif; ')
    html_content.append('  margin: 10px; ')
    html_content.append('  line-height: 1.4; ')
    html_content.append('  font-size: 12px; ')  # Уменьшен базовый размер шрифта
    html_content.append('  color: #333; ')
    html_content.append('  background-color: #fff; ')
    html_content.append('}')
    html_content.append('.table-container { ')
    html_content.append('  width: 100%; ')
    html_content.append('  overflow-x: auto; ')  # Горизонтальная прокрутка для таблиц
    html_content.append('  margin: 10px 0; ')
    html_content.append('  -webkit-overflow-scrolling: touch; ')  # Плавная прокрутка на iOS
    html_content.append('}')
    html_content.append('table { ')
    html_content.append('  border-collapse: collapse; ')
    html_content.append('  width: 100%; ')
    html_content.append('  min-width: 600px; ')  # Минимальная ширина для таблиц
    html_content.append('  font-size: 10px; ')  # Уменьшен размер шрифта в таблицах
    html_content.append('}')
    html_content.append('td, th { ')
    html_content.append('  border: 1px solid #ccc; ')
    html_content.append('  padding: 6px; ')  # Уменьшен padding
    html_content.append('  vertical-align: top; ')
    html_content.append('  word-wrap: break-word; ')
    html_content.append('  max-width: 200px; ')  # Ограничение ширины ячеек
    html_content.append('}')
    html_content.append('th { ')
    html_content.append('  background-color: #f8f9fa; ')
    html_content.append('  font-weight: bold; ')
    html_content.append('}')
    html_content.append('.first-row { ')  # Стиль для первой строки
    html_content.append('  font-size: 9px !important; ')  # Уменьшенный размер шрифта
    html_content.append('  padding: 4px 6px !important; ')  # Уменьшенный padding
    html_content.append('}')
    html_content.append('.first-col { ')  # Стиль для первого столбца
    html_content.append('  font-size: 9px !important; ')  # Уменьшенный размер шрифта
    html_content.append('  padding: 4px 6px !important; ')  # Уменьшенный padding
    html_content.append('}')
    html_content.append('.first-row.first-col { ')  # Стиль для ячейки в первой строке и первом столбце
    html_content.append('  font-size: 8px !important; ')  # Еще более уменьшенный размер
    html_content.append('  padding: 3px 5px !important; ')  # Еще более уменьшенный padding
    html_content.append('}')
    html_content.append('p { ')
    html_content.append('  margin: 6px 0; ')  # Уменьшен margin
    html_content.append('  font-size: 12px; ')  # Уменьшен размер шрифта
    html_content.append('}')
    html_content.append('h1 { font-size: 18px; margin: 12px 0 6px 0; }')
    html_content.append('h2 { font-size: 16px; margin: 10px 0 5px 0; }')
    html_content.append('h3 { font-size: 14px; margin: 8px 0 4px 0; }')
    html_content.append('h4 { font-size: 13px; margin: 7px 0 3px 0; }')
    html_content.append('h5 { font-size: 12px; margin: 6px 0 3px 0; }')
    html_content.append('h6 { font-size: 11px; margin: 5px 0 2px 0; }')
    html_content.append('@media (max-width: 480px) { ')
    html_content.append('  body { margin: 5px; font-size: 11px; }')
    html_content.append('  p { font-size: 11px; margin: 4px 0; }')
    html_content.append('  td, th { padding: 4px; font-size: 9px; }')
    html_content.append('  .first-row { font-size: 8px !important; padding: 3px 4px !important; }')
    html_content.append('  .first-col { font-size: 8px !important; padding: 3px 4px !important; }')
    html_content.append('  .first-row.first-col { font-size: 7px !important; padding: 2px 3px !important; }')
    html_content.append('  h1 { font-size: 16px; }')
    html_content.append('  h2 { font-size: 14px; }')
    html_content.append('  h3 { font-size: 13px; }')
    html_content.append('  h4 { font-size: 12px; }')
    html_content.append('  h5 { font-size: 11px; }')
    html_content.append('  h6 { font-size: 10px; }')
    html_content.append('}')
    html_content.append('@media (max-width: 360px) { ')
    html_content.append('  body { margin: 3px; font-size: 10px; }')
    html_content.append('  p { font-size: 10px; margin: 3px 0; }')
    html_content.append('  td, th { padding: 3px; font-size: 8px; }')
    html_content.append('  .first-row { font-size: 7px !important; padding: 2px 3px !important; }')
    html_content.append('  .first-col { font-size: 7px !important; padding: 2px 3px !important; }')
    html_content.append('  .first-row.first-col { font-size: 6px !important; padding: 1px 2px !important; }')
    html_content.append('  h1 { font-size: 14px; }')
    html_content.append('  h2 { font-size: 13px; }')
    html_content.append('  h3 { font-size: 12px; }')
    html_content.append('  h4 { font-size: 11px; }')
    html_content.append('  h5 { font-size: 10px; }')
    html_content.append('  h6 { font-size: 9px; }')
    html_content.append('}')
    html_content.append('</style>')
    html_content.append('</head>')
    html_content.append('<body>')
    
    # Обрабатываем все элементы документа в правильном порядке
    for element in doc.element.body:
        # Параграфы
        if element.tag.endswith('p'):
            for paragraph in doc.paragraphs:
                if paragraph._p == element:
                    if paragraph.text.strip() or paragraph.runs:
                        html_content.append(process_paragraph(paragraph))
                    break
        
        # Таблицы
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._tbl == element:
                    html_content.append(process_table(table))
                    break
        
        # Разрывы разделов (добавляем горизонтальную линию)
        elif element.tag.endswith('sectPr'):
            html_content.append('<hr style="margin: 15px 0; border: 0; border-top: 1px solid #ccc;">')
    
    html_content.append('</body>')
    html_content.append('</html>')
    
    # Сохраняем HTML файл
    try:
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_content))
        print(f"Файл успешно преобразован: {html_path}")
        return True
    except Exception as e:
        print(f"Ошибка при сохранении файла: {e}")
        return False

def main():
    if len(sys.argv) != 2:
        print("Использование: python docx_to_html.py <имя_файла.docx>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not docx_file.endswith('.docx'):
        print("Файл должен иметь расширение .docx")
        sys.exit(1)
    
    html_file = docx_file.replace('.docx', '.html')
    
    if docx_to_html(docx_file, html_file):
        print("Преобразование завершено успешно!")
    else:
        print("Произошла ошибка при преобразовании")
        sys.exit(1)

if __name__ == "__main__":
    main()
