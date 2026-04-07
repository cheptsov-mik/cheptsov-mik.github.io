#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Разбивает исходный .docx файл на несколько файлов, каждый из
которых содержит одну таблицу вместе с её надписью (перед таблицей)
и подписью (после таблицы). Параметры страницы (размер, ориентация,
поля) и форматирование текста (цвет, размер, полужирный, курсив,
подчёркивание и т.д.) копируются из оригинального документа.

Использование:
    python split_docx_by_tables.py исходный_файл.docx

Требования:
    pip install python-docx
"""

import sys
import copy
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
import t1

def copy_runs(src_paragraph, dst_paragraph):
    """Копирует все runs из src_paragraph в dst_paragraph, включая
    форматирование (цвет, размер, полужирный, курсив, подчёркивание и т.д.)."""
    for src_run in src_paragraph.runs:
        # Создаём новый run в целевом параграфе
        dst_run = dst_paragraph.add_run()
        dst_run.text = src_run.text

        # Копируем форматирование
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        dst_run.font.size = src_run.font.size
        dst_run.font.color.rgb = src_run.font.color.rgb
        dst_run.font.name = src_run.font.name
        dst_run.font.highlight_color = src_run.font.highlight_color
        dst_run.font.strike = src_run.font.strike
        dst_run.font.subscript = src_run.font.subscript
        dst_run.font.superscript = src_run.font.superscript

def copy_paragraph(paragraph):
    """Создаёт новый параграф, копируя текст и форматирование из
    исходного параграфа."""
    new_para = Document().add_paragraph()
    # Копируем стиль параграфа
    new_para.style = paragraph.style
    # Копируем все runs
    copy_runs(paragraph, new_para)
    return new_para

def copy_section_properties(src_section, dst_section):
    """Копирует настройки страницы (размер, ориентация, поля) из
    src_section в dst_section."""
    dst_section.page_width = src_section.page_width
    dst_section.page_height = src_section.page_height
    dst_section.orientation = src_section.orientation
    dst_section.top_margin = src_section.top_margin
    dst_section.bottom_margin = src_section.bottom_margin
    dst_section.left_margin = src_section.left_margin
    dst_section.right_margin = src_section.right_margin
    dst_section.header_distance = src_section.header_distance
    dst_section.footer_distance = src_section.footer_distance

def extract_caption_before(paragraphs, idx):
    """Возвращает параграф, который находится сразу перед таблицей,
    если он содержит текст. Иначе None."""
    if idx > 0:
        prev = paragraphs[idx - 1]
        if prev.text.strip():
            return prev
    return None

def extract_caption_after(paragraphs, idx):
    """Возвращает параграф, который находится сразу после таблицы,
    если он содержит текст. Иначе None."""
    if idx + 1 < len(paragraphs):
        nxt = paragraphs[idx + 1]
        if nxt.text.strip():
            return nxt
    return None

def split_document(doc_path):
    doc = Document(doc_path)

    # Список всех параграфов документа
    paragraphs = list(doc.paragraphs)

    # Список (таблица, надпись, подпись)
    tables_info = []

    # Текущий индекс параграфа в списке paragraphs
    para_idx = 0

    for table in doc.tables:
        # Пропускаем пустые параграфы, чтобы найти позицию таблицы
        while para_idx < len(paragraphs) and paragraphs[para_idx].runs == []:
            para_idx += 1

        caption_above = extract_caption_before(paragraphs, para_idx)
        caption_below = extract_caption_after(paragraphs, para_idx)

        tables_info.append((table, caption_above, caption_below))

        # Перейти к следующему параграфу после таблицы
        para_idx += 1

    # Получаем свойства страницы из первой секции оригинального документа
    original_section = doc.sections[0]

    # Создаём отдельные файлы
    for i, (table, cap_above, cap_below) in enumerate(tables_info, start=1):
        new_doc = Document()

        # Копируем настройки страницы
        copy_section_properties(original_section, new_doc.sections[0])

        # Добавляем надпись, если она есть
#        if cap_above:
#            new_para = copy_paragraph(cap_above)
#            new_doc._body._element.append(new_para._p)

        # Добавляем таблицу, копируя её XML‑элемент
        tbl_xml = copy.deepcopy(table._tbl)
        new_doc._body._element.append(tbl_xml)

        # Добавляем подпись, если она есть
#        if cap_below:
#            new_para = copy_paragraph(cap_below)
#            new_doc._body._element.append(new_para._p)

        # Формируем имя выходного файла
        base = Path(doc_path).stem
        out_name = f"{base}_table_{i}.docx"
        new_doc.save(out_name)
        print(f"Сохранено: {out_name}")

        html_name = f"{base}_table_{i}.html"
        t1.converter(out_name,html_name)


def main():
    if len(sys.argv) != 2:
        print("Использование: python split_docx_by_tables.py исходный_файл.docx")
        sys.exit(1)

    doc_path = sys.argv[1]
    if not Path(doc_path).is_file():
        print(f"Файл не найден: {doc_path}")
        sys.exit(1)

    split_document(doc_path)

if __name__ == "__main__":
    main()
